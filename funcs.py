from bs4 import BeautifulSoup as bs

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt

from requests_html import HTMLSession
from youtube_transcript_api import YouTubeTranscriptApi

import os


def assemble_str(transcr):

    result_str = str()
    for item in transcr:
        result_str += (item + ' ')
    return result_str

def write_to_doc(filename, write_str, combo):

    document = Document()

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Arial'

    p = document.add_paragraph('')
    p.add_run(filename.split('.')[0]).bold = True
    p.add_run('\n\n\n')

    if combo:

        blocks = write_str.split('\n\n')
        for block in blocks:
            strings = block.split('\n')
            if len(strings) == 2:
                p.add_run(strings[0]).bold = False
                p.add_run('\n')
                p.add_run(strings[1]).italic = True
                p.add_run('\n\n')

        document.save(filename)

    else:

        p = document.add_paragraph(write_str)
        document.save(filename)

def scripts_to_doc(script_strings, doc_fnames):

    curr_dir = os.getcwd()

    for script_string in script_strings:
        i = script_strings.index(script_string)
        if script_string:
            dir_name = doc_fnames[i].split('transcript')[0]
            if not os.path.exists(dir_name):
                os.makedirs(dir_name)

            file_path = os.path.join(curr_dir, dir_name, doc_fnames[i])
            write_to_doc(file_path, script_strings[i], i==2)

def clean_script(script):

    return " ".join(assemble_str(script).split("\n")).replace("  ", " ").replace(". ", ".\n\n")




def get_video_title(url):

    session = HTMLSession()
    response = session.get(url)
    response.html.render(sleep=1)
    soup = bs(response.html.html, "html.parser")
    video_title = soup.find("h1").text
    return video_title

def clean_filename(filename):
    
    return filename[1:].replace("\n", "").replace("?", "")

def char_cutoff_lines(script_words, char_lim):

    """

    this function assembles lines with a character 
    count cut-off per line

    """

    curr_line = str()
    lines = list()

    for i in range(len(script_words)):

        if len(curr_line) < char_lim:
            curr_line += script_words[i] + ' '
        else:
            lines.append(curr_line)
            curr_line = script_words[i] + ' '

    return lines

def txt_to_other(txt_fnames, other_format):

    return [" ".join(txt_fname.split('.')[:-1]) + other_format for txt_fname in txt_fnames]

def get_txt_fnames(vid_title, lang_strings):

    return [clean_filename(vid_title + f' transcript {lang_string}.txt') for lang_string in lang_strings]

def script_to_lines(script_str, char_lim):

    return char_cutoff_lines(script_str.split(' '), char_lim)



def text_to_pdf(txt_filename, pdf_filename):

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=15)

    f = open(txt_filename, "r")
    for x in f:
        pdf.cell(200,10, txt = x, ln = 1, align = 'C')

    pdf.output(pdf_filename)